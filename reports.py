import pandas as pd
from io import BytesIO
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor
import matplotlib.pyplot as plt
import tempfile
import os

# Watermark & Footer Text
DEDICATION_TEXT = (
    "To the Soul of my late father Abdulrahman, "
    "my long living mother Zainab, my beloved wife Yousra and my eyes, Abdulrahman & Osman."
)
FOOTER_TEXT = "© All Rights Reserved: Mohammedelnagi Mohammed"

def generate_excel(df: pd.DataFrame) -> BytesIO:
    output = BytesIO()
    with pd.ExcelWriter(output, engine="xlsxwriter") as writer:
        df.to_excel(writer, index=False, sheet_name="MalariaData")

        workbook = writer.book
        worksheet = writer.sheets["MalariaData"]

        # Optional: Style headers
        header_format = workbook.add_format({
            "bold": True,
            "text_wrap": True,
            "valign": "top",
            "fg_color": "#D7E4BC",
            "border": 1
        })
        for col_num, value in enumerate(df.columns.values):
            worksheet.write(0, col_num, value, header_format)

        # Optional: Set column width
        for i, width in enumerate(get_col_widths(df)):
            worksheet.set_column(i, i, width)

    output.seek(0)
    return output

def get_col_widths(df):
    col_widths = []
    for col in df.columns:
        max_len = max(
            df[col].astype(str).map(len).max(),
            len(col)
        )
        col_widths.append(min(max_len + 2, 50))  # max width limit
    return col_widths

def generate_word_report(df: pd.DataFrame) -> BytesIO:
    doc = Document()

    # Title
    title = doc.add_heading("Malaria Investigation Summary", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Dedication
    para = doc.add_paragraph()
    run = para.add_run(DEDICATION_TEXT)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0, 102, 204)  # blue-ish
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Summary stats
    doc.add_heading("Data Summary", level=1)
    doc.add_paragraph(f"Total Records: {len(df)}")

    # Charts
    chart_path = create_pie_chart(df)
    if chart_path:
        doc.add_picture(chart_path, width=None)

    # Table
    doc.add_heading("Data Table", level=1)
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = 'Table Grid'

    # Add headers
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(df.columns):
        hdr_cells[i].text = str(col)

    # Add data
    for index, row in df.iterrows():
        row_cells = table.add_row().cells
        for i, col in enumerate(df.columns):
            row_cells[i].text = str(row[col])

    # Footer (simulate watermark style)
    footer = doc.sections[0].footer
    footer_para = footer.paragraphs[0]
    footer_para.text = FOOTER_TEXT
    footer_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    footer_para.runs[0].italic = True
    footer_para.runs[0].font.size = Pt(9)

    # Save
    output = BytesIO()
    doc.save(output)
    output.seek(0)

    # Cleanup temp chart
    if chart_path and os.path.exists(chart_path):
        os.remove(chart_path)

    return output

def create_pie_chart(df: pd.DataFrame) -> str:
    """Creates a pie chart of gender distribution, returns path to image file"""
    if "Gender" not in df.columns:
        return ""

    counts = df["Gender"].value_counts()
    if counts.empty:
        return ""

    fig, ax = plt.subplots()
    ax.pie(counts, labels=counts.index, autopct="%1.1f%%", startangle=90)
    ax.set_title("Gender Distribution")
    plt.axis("equal")

    # Save to temp file
    tmpfile = tempfile.NamedTemporaryFile(delete=False, suffix=".png")
    plt.savefig(tmpfile.name)
    plt.close()

    return tmpfile.name