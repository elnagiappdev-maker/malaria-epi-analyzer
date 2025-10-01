import os
os.environ["STREAMLIT_WATCHER_TYPE"] = "poll"  # avoid inotify errors

import io, json
from datetime import datetime, date
import pandas as pd
import streamlit as st
import requests
import fitz  # PyMuPDF
from PIL import Image
from docx import Document
from docx.shared import Inches
import matplotlib.pyplot as plt

# ========== Configuration / Constants ==========
APP_NAME = "Malaria Epidemiological Analyzer"
DEDICATION_TEXT = "To the Soul of my Late Father, My Long Living Mother, My Lovely wife and My Eyes Abdulrahman & Osman"

GROUP_SIZE = 4  # pages per patient form
API_URL = "https://api.ocr.space/parse/image"

# ========== Authentication / Secrets ==========
try:
    VALID_USERNAME = st.secrets["APP_USERNAME"]
    VALID_PASSWORD = st.secrets["APP_PASSWORD"]
except KeyError:
    st.error("❌ Credentials not set. Add APP_USERNAME & APP_PASSWORD to Streamlit secrets.")
    st.stop()

# ========== Layout & Styles ==========
st.set_page_config(page_title=APP_NAME, layout="wide")
st.markdown("""
<style>
html, body, .stApp, [data-testid="stAppViewContainer"] {
  background: linear-gradient(180deg,#f7fafc 0%, #ffffff 80%);
}
.login-wrap { min-height: calc(100vh - 90px); display:flex; align-items:center; justify-content:center; padding:24px 16px; }
.login-card {
  width:100%; max-width:580px;
  background:#ffffff; border:2px solid #1f6feb; border-radius:16px;
  box-shadow:0 8px 24px rgba(0,0,0,0.12); padding:22px 20px;
}
.app-name { text-align:center; color:#1f6feb; font-weight:800; margin:6px 0 10px 0; font-size: clamp(20px, 4vw, 30px); }
.login-sub { text-align:center; color:#334155; margin-bottom:16px; font-size: clamp(12px, 2.5vw, 15px); }
.app-footer {
  position: fixed;
  left: 0; right: 0;
  bottom: 30px; /* raise above watermark */
  width: 100%;
  text-align: center;
  padding: 8px 8px 0 8px;
  background: #ffffffea;
  border-top: 1px solid #e5e7eb;
  z-index: 9999;
}
.footer-rights { font-weight:700; color:#1f6feb; margin-bottom:6px; }
.footer-dedication { font-weight:700; color:#0f172a; margin-bottom:8px; }
@media (max-width: 480px) { .login-card { padding:16px 14px; border-radius:14px; } }
</style>
""", unsafe_allow_html=True)

def render_footer():
    st.markdown(
        '<div class="app-footer">'
        '<div class="footer-rights">All Rights Reserved to Dr. Mohammedelnagi Mohammed</div>'
        f'<div class="footer-dedication"><b>Dedication:</b> {DEDICATION_TEXT}</div>'
        '</div>',
        unsafe_allow_html=True
    )

# ========== OCR / PDF Helpers ==========

def pdf_to_images(pdf_bytes: bytes):
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    images = []
    for pg in doc:
        pix = pg.get_pixmap(dpi=200)
        img = Image.open(io.BytesIO(pix.tobytes("png")))
        images.append(img)
    return images

def _post_ocr_image(image: Image.Image, filename: str, api_key: str, language: str = "eng,ara"):
    buf = io.BytesIO()
    image.save(buf, format="PNG")
    buf.seek(0)
    files = {"file": (filename, buf)}
    data = {
        "language": language,
        "isOverlayRequired": False,
        "OCREngine": 2,
        "scale": True,
        "isTable": True
    }
    headers = {"apikey": api_key}
    resp = requests.post(API_URL, data=data, files=files, headers=headers, timeout=300)
    resp.raise_for_status()
    j = resp.json()
    if j.get("IsErroredOnProcessing"):
        error_msg = j.get("ErrorMessage") or j.get("ErrorDetails") or "Unknown OCR error"
        raise RuntimeError(f"OCR Error: {error_msg}")
    return j["ParsedResults"][0]["ParsedText"]

def ocr_from_file_grouped(uploaded_file, api_key, languages="eng,ara", group_size=GROUP_SIZE):
    if hasattr(uploaded_file, "seek"):
        uploaded_file.seek(0)
    file_bytes = uploaded_file.read() if hasattr(uploaded_file, "read") else uploaded_file
    fname = getattr(uploaded_file, "name", "upload").lower()
    results = []
    if fname.endswith(".pdf"):
        imgs = pdf_to_images(file_bytes)
        for i in range(0, len(imgs), group_size):
            chunk = imgs[i:i+group_size]
            texts = []
            for j, img in enumerate(chunk):
                try:
                    t = _post_ocr_image(img, f"page_{i+j+1}.png", api_key, languages)
                    texts.append(t)
                except Exception as e:
                    texts.append(f"[ERROR on page {i+j+1}]: {e}")
            merged = "\n\n".join(texts)
            results.append(merged)
    else:
        img = Image.open(io.BytesIO(file_bytes))
        texts = _post_ocr_image(img, fname, api_key, languages)
        results.append(texts)
    return results

# ========== Parsing, Validation, etc. ==========

def infer_species(text: str):
    t = text.lower()
    for k in ["falciparum", "vivax", "ovale", "malariae", "knowlesi"]:
        if k in t:
            return k.capitalize()
    return "Unknown"

def normalize_record(rec: dict):
    # Dummy pass-through or apply your real normalizing logic
    return rec

def validate_record(rec: dict):
    # Dummy validation (replace with your logic)
    return {"valid": True}

# ========== Word Report with Charts ==========

def generate_word_report(df: pd.DataFrame, filename: str):
    doc = Document()
    doc.add_heading("Malaria Epidemiological Report", level=1)
    doc.add_paragraph(f"Date: {date.today().isoformat()}")
    # Add species pie chart
    if "species" in df.columns:
        species_counts = df["species"].fillna("Unknown").value_counts()
        fig, ax = plt.subplots()
        species_counts.plot.pie(autopct="%1.1f%%", ax=ax, legend=False)
        ax.set_ylabel("")
        fig.savefig("species_pie.png")
        doc.add_heading("Species Distribution", level=2)
        doc.add_picture("species_pie.png", width=Inches(4))
    # Add EpiCurve (cases over time)
    if "symptom_start" in df.columns:
        df["symptom_start"] = pd.to_datetime(df["symptom_start"], errors="coerce")
        df2 = df.dropna(subset=["symptom_start"])
        if not df2.empty:
            counts = df2.groupby(df2["symptom_start"].dt.date).size()
            fig2, ax2 = plt.subplots(figsize=(5, 3))
            counts.plot(kind="line", marker="o", ax=ax2)
            ax2.set_xlabel("Date")
            ax2.set_ylabel("Number of Cases")
            fig2.savefig("epicurve.png")
            doc.add_heading("EpiCurve", level=2)
            doc.add_picture("epicurve.png", width=Inches(5))
    # Add table of data (first few rows for simplicity)
    doc.add_heading("Summary Table", level=2)
    table = doc.add_table(rows=1, cols=len(df.columns))
    hdr = table.rows[0].cells
    for i, c in enumerate(df.columns):
        hdr[i].text = str(c)
    for _, row in df.iterrows():
        r_cells = table.add_row().cells
        for i, c in enumerate(df.columns):
            r_cells[i].text = str(row[c])
    out_path = filename
    doc.save(out_path)
    return out_path

# ========== Streamlit App Logic ==========

def render_login():
    st.markdown('<div class="login-wrap"><div class="login-card">', unsafe_allow_html=True)
    st.markdown(f'<div class="app-name">{APP_NAME}</div>', unsafe_allow_html=True)
    st.markdown('<div class="login-sub">Please enter your credentials.</div>', unsafe_allow_html=True)
    with st.form("login_form"):
        u = st.text_input("Username", key="login_user")
        p = st.text_input("Password", type="password", key="login_pass")
        submitted = st.form_submit_button("Log in")
        if submitted:
            if u == VALID_USERNAME and p == VALID_PASSWORD:
                st.session_state.logged_in = True
                st.rerun()
            else:
                st.error("Invalid username or password.")
    st.markdown('</div></div>', unsafe_allow_html=True)
    render_footer()

def render_app():
    st.title(APP_NAME)
    st.caption("OCR, process, and analyze malaria forms.")
    with st.sidebar:
        st.header("Account")
        st.success("Logged in")
        if st.button("Log out"):
            st.session_state.logged_in = False
            st.rerun()
        st.divider()
        st.header("Diagnostics")
        key = st.secrets.get("OCRSPACE_API_KEY", "")
        st.code(f"OCR key: {'configured' if key else 'Missing'}", language="bash")

    mode = st.radio("Mode:", ["Upload PDFs/Images", "Upload Excel"])
    records = []
    df = pd.DataFrame()

    if mode.startswith("Upload"):
        files = st.file_uploader("Upload files", type=["pdf","png","jpg","jpeg"], accept_multiple_files=True)
        ocr_key = st.secrets.get("OCRSPACE_API_KEY", "")
        if files and ocr_key:
            for f in files:
                with st.spinner(f"Processing {f.name}..."):
                    try:
                        form_texts = ocr_from_file_grouped(f, ocr_key)
                        for txt in form_texts:
                            rec = normalize_record({"raw_text": txt})
                            rec["species"] = infer_species(txt)
                            rec["symptom_start"] = None
                            rec["ingest_source"] = f.name
                            rec["created_at"] = datetime.utcnow().isoformat()
                            rec["validation_flags"] = json.dumps(validate_record(rec))
                            records.append(rec)
                    except Exception as e:
                        st.error(f"Error {f.name}: {e}")
            if records:
                df = pd.DataFrame(records)
                st.success(f"Processed {len(df)} record(s).")
                st.dataframe(df, use_container_width=True)
    else:
        xfile = st.file_uploader("Upload Excel (.xlsx)", type=["xlsx"])
        if xfile:
            try:
                df = pd.read_excel(xfile)
                for col in ["species", "symptom_start"]:
                    if col not in df.columns:
                        df[col] = None
                df["created_at"] = datetime.utcnow().isoformat()
                df["validation_flags"] = [json.dumps(validate_record(r.to_dict())) for _, r in df.iterrows()]
                st.success(f"Loaded {len(df)} rows from Excel.")
                st.dataframe(df, use_container_width=True)
            except Exception as e:
                st.error(f"Excel error: {e}")

    if not df.empty:
        st.subheader("Exports")
        if st.button("Download Excel"):
            st.download_button("Excel", df.to_excel(index=False), file_name="malaria_data.xlsx", mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
        report_name = st.text_input("Word report file name", value=f"report_{date.today().isoformat()}.docx")
        if st.button("Generate Word Report"):
            try:
                path = generate_word_report(df, report_name)
                with open(path, "rb") as fp:
                    st.download_button("Download Word Report", fp.read(), file_name=report_name)
            except Exception as e:
                st.error(f"Word report failed: {e}")

    render_footer()

def ensure_and_run():
    if "logged_in" not in st.session_state:
        st.session_state.logged_in = False
    if not st.session_state.logged_in:
        render_login()
    else:
        render_app()

if __name__ == "__main__":
    ensure_and_run()
